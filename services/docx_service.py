import io
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        sections: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    sections.append(" | ".join(row_cells))

        result = "\n".join(sections).strip()

        if not result:
            raise ValueError(
                "No extractable text found in the DOCX file. "
                "The document may be empty or image-based."
            )

        return result

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Failed to read DOCX file: {e}") from e